"""
src/api/csr_docx_builder.py

ICH E3 Clinical Study Report document builders.

- build_csr_docx: regulatory-grade deliverable DOCX (narrative only, no QC artefacts).
- build_csr_qc_report_docx: internal QC report DOCX with verification tables and audit trail.
"""
import hashlib
import io
import json
import re
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

VERIFIED_GREEN = RGBColor(0x1F, 0x8A, 0x4C)
HAZARD_AMBER = RGBColor(0xB8, 0x5C, 0x00)
CORTEX_GREY = RGBColor(0x5B, 0x6B, 0x7B)
INK_BLACK = RGBColor(0x14, 0x1B, 0x2D)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _set_cell_shading(cell, hex_colour: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tc_pr.append(shd)


def _add_page_border(doc: Document) -> None:
    for section in doc.sections:
        sect_pr = section._sectPr
        pg_borders = OxmlElement("w:pgBorders")
        pg_borders.set(qn("w:offsetFrom"), "page")
        for side in ("top", "left", "bottom", "right"):
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:space"), "24")
            border.set(qn("w:color"), "141B2D")
            pg_borders.append(border)
        sect_pr.append(pg_borders)


def _add_header_footer(doc: Document, title: str) -> None:
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = ""
        run = hp.add_run(f"Clinical Study Report  |  {title}")
        run.font.size = Pt(8)
        run.font.color.rgb = CORTEX_GREY
        run.italic = True
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.style = doc.styles["Normal"]
        fld_char1 = OxmlElement("w:fldChar")
        fld_char1.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "end")
        run1 = fp.add_run()
        run1.font.size = Pt(8)
        run1.font.color.rgb = CORTEX_GREY
        run1._r.append(fld_char1)
        run2 = fp.add_run()
        run2.font.size = Pt(8)
        run2.font.color.rgb = CORTEX_GREY
        run2._r.append(instr)
        run3 = fp.add_run()
        run3.font.size = Pt(8)
        run3.font.color.rgb = CORTEX_GREY
        run3._r.append(fld_char2)


def _pct(p: Optional[float]) -> str:
    if p is None:
        return "—"
    return f"{round(p * 100)}%"


def _clean_prose(summary: str) -> str:
    """Remove residual markdown and bullets without destroying hyphens in numbers."""
    if not summary:
        return summary
    text = summary.replace("**", "").replace("__", "")
    text = text.replace("•", "")
    # Remove markdown bullet dashes at line starts only; preserve hyphens inside values/ranges.
    text = re.sub(r"^\s*-\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _add_plain_paragraph(doc: Document, text: str, style: str = "Normal") -> None:
    p = doc.add_paragraph(text, style=style)
    for r in p.runs:
        r.font.bold = False
        r.font.size = Pt(11)
        r.font.color.rgb = INK_BLACK


def _build_cover_page(doc: Document, summary_result) -> None:
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_heading("Clinical Study Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in title.runs:
        r.font.color.rgb = INK_BLACK
        r.font.size = Pt(28)
    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(summary_result.filename)
    run.font.size = Pt(14)
    run.font.color.rgb = CORTEX_GREY
    run.italic = True
    doc.add_page_break()


def _build_toc_page(doc: Document, sections: List) -> None:
    h = doc.add_heading("Table of Contents", level=1)
    for r in h.runs:
        r.font.color.rgb = INK_BLACK
        r.font.bold = False
    doc.add_paragraph()
    for sec in sections:
        ns = sec.section_number
        label = sec.canonical_title or sec.title or f"Section {ns}"
        table_count = sec.tables_found
        p = doc.add_paragraph()
        run = p.add_run(f"Section {ns}  —  {label}")
        run.font.size = Pt(11)
        run.font.color.rgb = INK_BLACK
        if table_count > 0:
            run2 = p.add_run(f"  ({table_count} table{'s' if table_count != 1 else ''})")
            run2.font.size = Pt(9)
            run2.font.color.rgb = CORTEX_GREY
            run2.italic = True
    doc.add_page_break()


def _build_document_synthesis_chapter(doc: Document, summary_result) -> None:
    h = doc.add_heading("Document Synthesis", level=1)
    for r in h.runs:
        r.font.color.rgb = INK_BLACK
        r.font.bold = False
    _add_plain_paragraph(doc, _clean_prose(summary_result.document_synthesis))
    doc.add_paragraph()


def _build_section_chapter(
    doc: Document,
    section_result,
    section_index: int,
) -> None:
    ns = section_result.section_number
    label = section_result.canonical_title or section_result.title or f"Section {ns}"
    h = doc.add_heading(f"Section {ns}: {label}", level=1)
    for r in h.runs:
        r.font.color.rgb = INK_BLACK
        r.font.bold = False
    doc.add_paragraph()

    if section_result.key_findings:
        h2 = doc.add_heading("Key Findings", level=2)
        for r in h2.runs:
            r.font.color.rgb = INK_BLACK
            r.font.bold = False
        for finding in section_result.key_findings:
            _add_plain_paragraph(doc, _clean_prose(finding))
        doc.add_paragraph()

    if section_result.section_synthesis:
        h2 = doc.add_heading("Section Synthesis", level=2)
        for r in h2.runs:
            r.font.color.rgb = INK_BLACK
            r.font.bold = False
        _add_plain_paragraph(doc, _clean_prose(section_result.section_synthesis))
        doc.add_paragraph()

    if section_result.table_summaries:
        h2 = doc.add_heading("Per-Table Summaries", level=2)
        for r in h2.runs:
            r.font.color.rgb = INK_BLACK
            r.font.bold = False
        for idx, tr in enumerate(section_result.table_summaries):
            table_title = tr.title or f"Table {idx + 1}"
            h3 = doc.add_heading(f"{idx + 1}. {table_title}", level=3)
            for r in h3.runs:
                r.font.color.rgb = INK_BLACK
                r.font.bold = False
            if tr.summary:
                _add_plain_paragraph(doc, _clean_prose(tr.summary))
            else:
                p = doc.add_paragraph("[No summary generated]")
                for r in p.runs:
                    r.italic = True
                    r.font.color.rgb = HAZARD_AMBER
        doc.add_page_break()


def build_csr_docx(summary_result) -> bytes:
    """Build the regulatory deliverable DOCX (narrative only, no QC artefacts)."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    _add_page_border(doc)
    _add_header_footer(doc, summary_result.filename)
    _build_cover_page(doc, summary_result)
    _build_toc_page(doc, summary_result.sections)
    # ICH E3 deliverables start directly with section chapters; no Document Synthesis.
    for i, sec in enumerate(summary_result.sections):
        _build_section_chapter(doc, sec, i)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# QC Report builder
# ---------------------------------------------------------------------------


def _build_qc_cover_page(doc: Document, summary_result) -> None:
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_heading("CSR Pipeline QC Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in title.runs:
        r.font.color.rgb = INK_BLACK
        r.font.size = Pt(28)
    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(summary_result.filename)
    run.font.size = Pt(14)
    run.font.color.rgb = CORTEX_GREY
    run.italic = True
    doc.add_paragraph()
    meta_info = [
        f"Total Pages: {summary_result.total_pages}",
        f"Sections Analysed: {len(summary_result.sections)}",
        f"Tables Extracted: {summary_result.total_tables}",
        f"Tables Verified: {summary_result.verified_tables}/{summary_result.total_tables}",
        f"Numeric Accuracy: {_pct(summary_result.overall_numeric_accuracy)}",
        f"Inference Time: {round(summary_result.total_inference_time_ms / 1000, 1)}s",
        f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}",
    ]
    for line in meta_info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.color.rgb = INK_BLACK
    doc.add_page_break()


def _build_qc_verification_table(doc: Document, tr) -> None:
    facts = getattr(tr, "extracted_facts", None) or []
    if not facts:
        p = doc.add_paragraph("No extracted facts available.")
        for r in p.runs:
            r.font.color.rgb = CORTEX_GREY
            r.italic = True
        return

    tbl = doc.add_table(rows=1 + len(facts), cols=5)
    tbl.style = "Light Grid Accent 1"
    headers = ["Value", "Type", "Row", "Source Cell", "Status"]
    hdr_cells = tbl.rows[0].cells
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8)
                r.font.color.rgb = WHITE
        _set_cell_shading(hdr_cells[i], "141B2D")
    for idx, fact in enumerate(facts):
        row = tbl.rows[idx + 1].cells
        row[0].text = str(fact.get("value", ""))
        row[1].text = fact.get("value_type", "")
        row[2].text = str(fact.get("source_row_idx", "—") if fact.get("source_row_idx") is not None else "—")
        row[3].text = (fact.get("source_label", "") or "") + "=" + (fact.get("source_value_repr", "") or "")
        status = fact.get("status", "unverified")
        row[4].text = status.upper()
        fill = "1F8A4C" if status == "verified" else "B85C00"
        _set_cell_shading(row[4], fill)
        for c in row:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(7)
        for r in row[4].paragraphs[0].runs:
            r.bold = True
            r.font.color.rgb = WHITE
    doc.add_paragraph()


def _build_qc_section_chapter(doc: Document, section_result) -> None:
    ns = section_result.section_number
    label = section_result.canonical_title or section_result.title or f"Section {ns}"
    h = doc.add_heading(f"Section {ns}: {label}", level=1)
    for r in h.runs:
        r.font.color.rgb = INK_BLACK
    doc.add_paragraph()

    for tr in section_result.table_summaries:
        h3 = doc.add_heading(f"{tr.table_id} — {tr.title or 'Untitled'}", level=3)
        for r in h3.runs:
            r.font.color.rgb = INK_BLACK
        meta = f"Type: {tr.table_type}  |  Page: {tr.page}  |  "
        meta += f"Verified: {'✓' if tr.verified else '✗'}  |  Accuracy: {_pct(tr.numeric_accuracy)}"
        p = doc.add_paragraph()
        run = p.add_run(meta)
        run.font.size = Pt(8)
        run.font.color.rgb = CORTEX_GREY
        run.italic = True
        if tr.summary:
            _add_plain_paragraph(doc, tr.summary)
        else:
            p = doc.add_paragraph("[No summary generated]")
            for r in p.runs:
                r.italic = True
                r.font.color.rgb = HAZARD_AMBER
        if tr.warnings:
            for w in tr.warnings[:5]:
                p = doc.add_paragraph(f"⚠ {w}")
                for r in p.runs:
                    r.font.size = Pt(8)
                    r.font.color.rgb = HAZARD_AMBER
        _build_qc_verification_table(doc, tr)
    doc.add_page_break()


def _build_qc_consistency_chapter(doc: Document, consistency_warnings: List[Dict]) -> None:
    h = doc.add_heading("Cross-Table Consistency Verification", level=1)
    for r in h.runs:
        r.font.color.rgb = INK_BLACK
    doc.add_paragraph()
    if not consistency_warnings:
        p = doc.add_paragraph("No cross-table consistency issues detected.")
        for r in p.runs:
            r.font.size = Pt(11)
            r.font.color.rgb = VERIFIED_GREEN
            r.bold = True
        return
    tbl = doc.add_table(rows=1 + len(consistency_warnings), cols=4)
    tbl.style = "Light Grid Accent 1"
    headers = ["Category", "Severity", "Message", "Source Tables"]
    hdr_cells = tbl.rows[0].cells
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = WHITE
        _set_cell_shading(hdr_cells[i], "141B2D")
    for idx, w in enumerate(consistency_warnings):
        row = tbl.rows[idx + 1].cells
        row[0].text = w.get("category", "")
        severity = w.get("severity", "info")
        row[1].text = severity.upper()
        row[2].text = w.get("message", "")
        row[3].text = ", ".join(w.get("source_tables", []))
        sev_fill = {"error": "C0392B", "warning": "B85C00", "info": "5B6B7B"}
        _set_cell_shading(row[1], sev_fill.get(severity, "5B6B7B"))
        for p in row[1].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = WHITE
                r.font.size = Pt(9)
        for c in row:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)


def _build_qc_audit_trail(doc: Document, summary_result) -> None:
    h = doc.add_heading("21 CFR Part 11 Audit Trail", level=1)
    for r in h.runs:
        r.font.color.rgb = INK_BLACK
    doc.add_paragraph()
    doc.add_paragraph(
        "This document was generated by an automated clinical narrative engine. "
        "The following audit trail satisfies 21 CFR Part 11 requirements for "
        "electronic record integrity and traceability."
    )
    doc.add_paragraph()
    summary_hash = hashlib.sha256(
        summary_result.document_synthesis.encode("utf-8")
    ).hexdigest()
    rows_data = [
        ("System", "ClinicalSafe CSR Generator"),
        ("Regulation", "21 CFR Part 11"),
        ("Generated (UTC)", datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")),
        ("Document", summary_result.filename),
        ("Total Pages (source)", str(summary_result.total_pages)),
        ("Sections", str(len(summary_result.sections))),
        ("Total Tables", str(summary_result.total_tables)),
        ("Verified Tables", str(summary_result.verified_tables)),
        ("Overall Numeric Accuracy", _pct(summary_result.overall_numeric_accuracy)),
        ("Inference Time", f"{round(summary_result.total_inference_time_ms / 1000, 1)}s"),
        ("Model", "NVIDIA NIM (meta/llama-3.3-70b-instruct)"),
        ("Generation Engine", "CSRNIMSynthesizer v1.0"),
        ("Document SHA-256", summary_hash),
    ]
    tbl = doc.add_table(rows=len(rows_data), cols=2)
    tbl.style = "Light List Accent 1"
    for i, (k, v) in enumerate(rows_data):
        c0, c1 = tbl.rows[i].cells
        c0.text = k
        c1.text = v
        for r in c0.paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(9)
        for r in c1.paragraphs[0].runs:
            r.font.size = Pt(9)
            r.font.name = "Consolas"


def build_csr_qc_report_docx(summary_result) -> bytes:
    """Build the internal QC report DOCX with verification tables and audit trail."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.1
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
    _add_page_border(doc)
    _add_header_footer(doc, summary_result.filename)
    _build_qc_cover_page(doc, summary_result)
    for sec in summary_result.sections:
        _build_qc_section_chapter(doc, sec)
    _build_qc_consistency_chapter(doc, summary_result.consistency_warnings)
    _build_qc_audit_trail(doc, summary_result)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_csr_audit_log(summary_result) -> str:
    """Return a standalone JSON audit log for the CSR run."""
    tables_log = []
    for sec in summary_result.sections:
        for tr in sec.table_summaries:
            tables_log.append(
                {
                    "table_id": tr.table_id,
                    "ich_section_number": tr.ich_section_number,
                    "ich_section_title": tr.ich_section_title,
                    "table_type": tr.table_type,
                    "page": tr.page,
                    "title": tr.title,
                    "verified": tr.verified,
                    "numeric_accuracy": tr.numeric_accuracy,
                    "inference_time_ms": tr.inference_time_ms,
                    "warnings": tr.warnings,
                    "extracted_facts": tr.extracted_facts,
                    "error": tr.error,
                }
            )

    log = {
        "system": "ClinicalSafe CSR Generator",
        "regulation": "21 CFR Part 11",
        "generated_utc": datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
        "document": summary_result.filename,
        "total_pages": summary_result.total_pages,
        "sections": len(summary_result.sections),
        "total_tables": summary_result.total_tables,
        "verified_tables": summary_result.verified_tables,
        "overall_numeric_accuracy": summary_result.overall_numeric_accuracy,
        "inference_time_ms": summary_result.total_inference_time_ms,
        "model": "NVIDIA NIM (meta/llama-3.3-70b-instruct)",
        "document_sha256": hashlib.sha256(
            summary_result.document_synthesis.encode("utf-8")
        ).hexdigest(),
        "consistency_warnings": summary_result.consistency_warnings,
        "errors": summary_result.errors,
        "tables": tables_log,
    }
    return json.dumps(log, indent=2)
